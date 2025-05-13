
import streamlit as st
import fitz
from supabase import create_client, Client
from datetime import datetime
from docx import Document
from io import BytesIO
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

SUPABASE_URL = "https://syrznbowqhvooxwzikhf.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InN5cnpuYm93cWh2b294d3ppa2hmIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDcwNjEzMDUsImV4cCI6MjA2MjYzNzMwNX0.IqeOV-3hynzr2mSN9quFlfkBEaqTKF6LwpL6IlmqYoU"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

secoes_lab = {
    "Bioquímica": ["glicose", "uréia", "creatinina", "potássio", "sódio", "cálcio", "fósforo"],
    "Hematologia": ["hemoglobina", "hematócrito", "vcm", "hcm", "leucócitos", "plaquetas"],
    "Hormônios": ["tsh", "t4", "t3", "pth"],
    "Vitaminas e Metabolismo Mineral": ["vitamina", "b12", "ácido fólico"],
    "Urina Tipo I": ["ph", "densidade", "hemácias", "leucócitos", "proteína"]
}

ruidos = ["cnpj", "crm", "laboratório", "assinatura", "referência", "nota", "método", "validado"]

def extrair_texto(pdf_file):
    texto = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page in doc:
            texto += page.get_text()
    return texto


def normalizar_linha(linha):
    linha = linha.lower().strip()
    linha = re.sub(r'\s+', ' ', linha)
    linha = linha.replace(":", "")
    return linha

def classificar_exames(texto):
    linhas = texto.lower().splitlines()
    dados = {secao: [] for secao in secoes_lab}
    dados["Outros"] = []
    vistos = set()

    for linha in linhas:
        l = linha.strip()
        if not l or any(p in l for p in ruidos): continue
        chave = normalizar_linha(l)
        if chave in vistos: continue
        vistos.add(chave)

        adicionou = False
        for secao, termos in secoes_lab.items():
            if any(t in chave for t in termos):
                dados[secao].append(l.capitalize())
                adicionou = True
                break
        if not adicionou and ":" in l:
            dados["Outros"].append(l.capitalize())
    return dados

    linhas = texto.lower().splitlines()
    dados = {secao: [] for secao in secoes_lab}
    dados["Outros"] = []
    for linha in linhas:
        l = linha.strip()
        if not l or any(p in l for p in ruidos): continue
        adicionou = False
        for secao, termos in secoes_lab.items():
            if any(t in l for t in termos):
                dados[secao].append(l.capitalize())
                adicionou = True
                break
        if not adicionou and ":" in l:
            dados["Outros"].append(l.capitalize())
    return dados

def limpar_texto(texto):
    texto = texto.encode("utf-8", "ignore").decode("utf-8", "ignore")
    texto = re.sub(r'[^ -~À-ÿ:\.,\-/()\[\] ]+', '', texto)
    return texto.strip()

def gerar_docx_laboratorial(nome, data, dados):
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Relatório de Exames Laboratoriais", 0)
    
    p1 = doc.add_paragraph()
    run1 = p1.add_run(f"Paciente: {nome}")
    run1.bold = True
    run1.font.size = Pt(11)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Data da coleta: {data}")
    run2.font.size = Pt(10)

    doc.add_paragraph("")

    for secao, itens in dados.items():
        if itens:
            doc.add_heading(secao, level=1)
            for item in itens:
                try:
                    texto_limpo = "".join(c for c in item if 32 <= ord(c) <= 126 or c in "\n\t .,:-_/()[]%")
                    doc.add_paragraph(texto_limpo.strip(), style="List Bullet")
                except Exception:
                    doc.add_paragraph("**Erro ao processar item**", style="List Bullet")
            doc.add_paragraph("-" * 40)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def gerar_docx_imagem(nome, data, texto):
    doc = Document()
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    run = header.add_run()
    run.add_picture("logo.png", width=Inches(2.5))
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Laudo de Imagem", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Paciente: {nome}")
    doc.add_paragraph(f"Data do exame: {data}")
    doc.add_paragraph("")

    for linha in texto.splitlines():
        texto_limpo = limpar_texto(linha)
        if texto_limpo:
            doc.add_paragraph(texto_limpo)

    rodape = section.footer.paragraphs[0]
    rodape.text = "Dr. João Batista Zinato – CRM 42099"
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape.runs[0].font.size = Pt(9)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()

st.title("Aplicativo de Laudos – Dr. João Batista Zinato")

with st.form("formulario"):
    tipo = st.radio("Tipo de exame", ["Exame Laboratorial", "Exame de Imagem"])
    nome = st.text_input("Nome do paciente")
    cpf = st.text_input("CPF")
    data = st.date_input("Data do exame")
    arquivo_pdf = st.file_uploader("Envie o PDF do laudo", type="pdf")
    enviar = st.form_submit_button("Processar")

if enviar and arquivo_pdf:
    texto = extrair_texto(arquivo_pdf)
    supabase.table("laudos").insert({
        "nome": nome, "cpf": cpf, "data_nasc": None,
        "data_laudo": datetime.now().isoformat(),
        "conteudo": texto
    }).execute()

    if tipo == "Exame Laboratorial":
        dados = classificar_exames(texto)
        docx_file = gerar_docx_laboratorial(nome, str(data), dados)
    else:
        docx_file = gerar_docx_imagem(nome, str(data), texto)

    st.success("Relatório gerado com sucesso.")
    st.download_button("Baixar relatório", docx_file, file_name="relatorio.docx")
